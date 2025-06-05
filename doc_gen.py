from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches


def render_cv_docx(cv, infos_perso, file_path):
    doc = Document()
    doc.add_heading(f"{infos_perso.get('prenom', '')} {infos_perso.get('nom', '')}", 0)
    doc.add_paragraph(
        f"{infos_perso.get('adresse', '')}\n"
        f"{infos_perso.get('telephone', '')} | "
        f"{infos_perso.get('email', '')} | {infos_perso.get('age', '')} ans"
    )
    if cv.get("profil"):
        doc.add_heading("Profil professionnel", level=1)
        doc.add_paragraph(cv.get("profil"))
    if cv.get("competences"):
        doc.add_heading("Compétences clés", level=1)
        for c in cv.get("competences"):
            doc.add_paragraph(c, style="List Bullet")
    if cv.get("experiences"):
        doc.add_heading("Expériences professionnelles", level=1)
        for e in cv.get("experiences"):
            doc.add_paragraph(e, style="List Bullet")
    if cv.get("formations"):
        doc.add_heading("Formations", level=1)
        for f in cv.get("formations"):
            doc.add_paragraph(f, style="List Bullet")
    if cv.get("autres"):
        doc.add_heading("Autres informations", level=1)
        for a in cv.get("autres"):
            doc.add_paragraph(a, style="List Bullet")
    doc.save(file_path)

def render_lm_docx(lettre_motivation, infos_perso, file_path):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    header_para = section.header.paragraphs[0]
    header_para.text = f"{infos_perso.get('prenom', '')} {infos_perso.get('nom', '')}"
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading("Lettre de motivation", 0)

    info_para = doc.add_paragraph(
        f"{infos_perso.get('adresse', '')}\n"
        f"{infos_perso.get('telephone', '')} | "
        f"{infos_perso.get('email', '')}"
    )
    info_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for block in lettre_motivation.split("\n\n"):
        p = doc.add_paragraph(block.strip())
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(12)

    sign = doc.add_paragraph(
        f"{infos_perso.get('prenom', '')} {infos_perso.get('nom', '')}"
    )
    sign.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    footer_para = section.footer.paragraphs[0]
    footer_para.text = f"{infos_perso.get('prenom', '')} {infos_perso.get('nom', '')}"
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(file_path)



def render_fiche_docx(fiche, file_path):
    doc = Document()
    doc.add_heading(fiche.get("titre", "Fiche de poste"), 0)
    doc.add_paragraph(f"Employeur : {fiche.get('employeur', '')}")
    doc.add_paragraph(f"Ville : {fiche.get('ville', '')}")
    doc.add_paragraph(f"Salaire : {fiche.get('salaire', '')}")
    doc.add_paragraph(f"Type de contrat : {fiche.get('type_contrat', '')}")
    if fiche.get("missions"):
        doc.add_heading("Missions principales", level=1)
        for m in fiche.get("missions", []):
            doc.add_paragraph(m, style="List Bullet")
    if fiche.get("competences"):
        doc.add_heading("Compétences requises", level=1)
        for c in fiche.get("competences", []):
            doc.add_paragraph(c, style="List Bullet")
    if fiche.get("savoir_etre"):
        doc.add_heading("Savoir-être", level=1)
        for s in fiche.get("savoir_etre", []):
            doc.add_paragraph(s, style="List Bullet")
    if fiche.get("avantages"):
        doc.add_heading("Avantages", level=1)
        for a in fiche.get("avantages", []):
            doc.add_paragraph(a, style="List Bullet")
    if fiche.get("autres"):
        doc.add_heading("Autres informations", level=1)
        for x in fiche.get("autres", []):
            doc.add_paragraph(x, style="List Bullet")
    doc.save(file_path)
