from flask import Flask, render_template, request, send_file, url_for, redirect
from jinja2 import Template
import pdfkit
import os
import platform
import shutil
from docx import Document
import re
import requests
import time
import tempfile
import json
import uuid

import pytesseract
from pdf2image import convert_from_path
from PyPDF2 import PdfReader

GROQ_API_KEY = "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"
GROQ_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"

TMP_DIR = "tmp"
os.makedirs(TMP_DIR, exist_ok=True)

app = Flask(__name__)

# --------- UTILS EXTRACTION ---------
def extract_text_from_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        fulltext = "\n".join([page.extract_text() or "" for page in reader.pages])
        if fulltext.strip():
            return fulltext
    except Exception:
        pass
    try:
        images = convert_from_path(file_path)
        text = "\n".join([pytesseract.image_to_string(img, lang='fra+eng') for img in images])
        return text
    except Exception:
        return ""

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception:
        return ""

# --------- IA / JSON ---------
def ask_groq(prompt):
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {GROQ_API_KEY}"
    }
    data = {
        "model": GROQ_MODEL,
        "messages": [
            {"role": "system", "content": "Tu es un assistant RH expert, spécialiste du recrutement en France."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.2,
        "max_tokens": 2800
    }
    resp = requests.post(url, headers=headers, json=data, timeout=80)
    j = resp.json()
    content = j["choices"][0]["message"]["content"]
    return content

def extract_first_json(text):
    m = re.search(r'(\{[\s\S]+\})', text)
    if not m:
        return None
    try:
        return json.loads(m.group(1))
    except Exception:
        text_clean = m.group(1).replace('\n', '').replace('\r', '')
        try:
            return json.loads(text_clean)
        except Exception:
            return None

def find_wkhtmltopdf():
    path = shutil.which("wkhtmltopdf")
    return path

if platform.system() == "Windows":
    WKHTMLTOPDF_PATH = r"C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
    config = pdfkit.configuration(wkhtmltopdf=WKHTMLTOPDF_PATH)
else:
    wkhtmltopdf_path = find_wkhtmltopdf()
    if wkhtmltopdf_path:
        config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
    else:
        raise RuntimeError("wkhtmltopdf non trouvé sur le système Railway ! Vérifie l'installation.")

# --------- DOCX/PDF GENERATION ---------
def render_cv_docx(cv, infos_perso, file_path):
    doc = Document()
    doc.add_heading(f"{infos_perso['prenom']} {infos_perso['nom']}", 0)
    doc.add_paragraph(f"{infos_perso['adresse']}\n{infos_perso['telephone']} | {infos_perso['email']} | {infos_perso['age']} ans")
    if cv.get('profil'):
        doc.add_heading("Profil professionnel", level=1)
        doc.add_paragraph(cv.get('profil'))
    if cv.get('competences'):
        doc.add_heading("Compétences clés", level=1)
        for c in cv.get('competences'):
            doc.add_paragraph(c, style='List Bullet')
    if cv.get('experiences'):
        doc.add_heading("Expériences professionnelles", level=1)
        for e in cv.get('experiences'):
            doc.add_paragraph(e, style='List Bullet')
    if cv.get('formations'):
        doc.add_heading("Formations", level=1)
        for f in cv.get('formations'):
            doc.add_paragraph(f, style='List Bullet')
    if cv.get('autres'):
        doc.add_heading("Autres informations", level=1)
        for a in cv.get('autres'):
            doc.add_paragraph(a, style='List Bullet')
    doc.save(file_path)

def render_lm_docx(lettre_motivation, infos_perso, file_path):
    doc = Document()
    doc.add_heading("Lettre de motivation", 0)
    doc.add_paragraph(f"{infos_perso['prenom']} {infos_perso['nom']}\n{infos_perso['adresse']}\n{infos_perso['telephone']} | {infos_perso['email']} | {infos_perso['age']} ans")
    doc.add_paragraph(lettre_motivation)
    doc.save(file_path)

def render_fiche_docx(fiche, file_path):
    doc = Document()
    doc.add_heading(fiche.get("titre","Fiche de poste"), 0)
    doc.add_paragraph(f"Employeur : {fiche.get('employeur','')}")
    doc.add_paragraph(f"Ville : {fiche.get('ville','')}")
    doc.add_paragraph(f"Salaire : {fiche.get('salaire','')}")
    doc.add_paragraph(f"Type de contrat : {fiche.get('type_contrat','')}")
    if fiche.get("missions"):
        doc.add_heading("Missions principales", level=1)
        for m in fiche.get("missions", []):
            doc.add_paragraph(m, style='List Bullet')
    if fiche.get("competences"):
        doc.add_heading("Compétences requises", level=1)
        for c in fiche.get("competences", []):
            doc.add_paragraph(c, style='List Bullet')
    if fiche.get("savoir_etre"):
        doc.add_heading("Savoir-être", level=1)
        for s in fiche.get("savoir_etre", []):
            doc.add_paragraph(s, style='List Bullet')
    if fiche.get("avantages"):
        doc.add_heading("Avantages", level=1)
        for a in fiche.get("avantages", []):
            doc.add_paragraph(a, style='List Bullet')
    if fiche.get("autres"):
        doc.add_heading("Autres informations", level=1)
        for x in fiche.get("autres", []):
            doc.add_paragraph(x, style='List Bullet')
    doc.save(file_path)

# --------- ROUTES ---------
@app.route('/', methods=['GET', 'POST'])
def index():
    error = ""
    results = {}

    # PATCH: Initialiser infos_perso dès le début pour TOUS les cas
    infos_perso = {
        "nom": "", "prenom": "", "adresse": "",
        "telephone": "", "email": "", "age": ""
    }

    context = {
        "nom": "", "prenom": "", "adresse": "", "telephone": "", "email": "", "age": "",
        "xp_poste": [], "xp_entreprise": [], "xp_lieu": [], "xp_debut": [], "xp_fin": [],
        "dip_titre": [], "dip_lieu": [], "dip_date": [],
        "description": "",
        "error": ""
    }

    if request.method == 'POST':
        nom = request.form.get('nom', '').strip()
        prenom = request.form.get('prenom', '').strip()
        adresse = request.form.get('adresse', '').strip()
        telephone = request.form.get('telephone', '').strip()
        email = request.form.get('email', '').strip()
        age = request.form.get('age', '').strip()
        description = request.form.get('description', '').strip()
        xp_poste = request.form.getlist('xp_poste')
        xp_entreprise = request.form.getlist('xp_entreprise')
        xp_lieu = request.form.getlist('xp_lieu')
        xp_debut = request.form.getlist('xp_debut')
        xp_fin = request.form.getlist('xp_fin')
        dip_titre = request.form.getlist('dip_titre')
        dip_lieu = request.form.getlist('dip_lieu')
        dip_date = request.form.getlist('dip_date')
        cv_file = request.files.get('cv_file')

        context.update({
            "nom": nom, "prenom": prenom, "adresse": adresse, "telephone": telephone, "email": email, "age": age,
            "xp_poste": xp_poste, "xp_entreprise": xp_entreprise, "xp_lieu": xp_lieu, "xp_debut": xp_debut, "xp_fin": xp_fin,
            "dip_titre": dip_titre, "dip_lieu": dip_lieu, "dip_date": dip_date,
            "description": description
        })

        cv_uploaded_text = ""
        if cv_file and cv_file.filename:
            ext = cv_file.filename.lower().split('.')[-1]
            with tempfile.NamedTemporaryFile(delete=False, suffix="." + ext) as tmp:
                cv_file.save(tmp.name)
                file_path = tmp.name
            if ext == "pdf":
                cv_uploaded_text = extract_text_from_pdf(file_path)
            elif ext == "docx":
                cv_uploaded_text = extract_text_from_docx(file_path)
            else:
                error = "Format de CV non supporté (PDF ou DOCX uniquement)"
            os.unlink(file_path)

        fiche_poste = {}
        file_id = uuid.uuid4().hex
        # -------------- IA ROUTINE -----------------
        if cv_uploaded_text.strip():
            prompt_parse_cv = f"""
Lis attentivement le texte suivant extrait d’un CV PDF ou DOCX. Trie les informations dans ce JSON, section par section, sans jamais inventer :

{{
  "nom": "...",
  "prenom": "...",
  "adresse": "...",
  "telephone": "...",
  "email": "...",
  "age": "...",
  "profil": "...",
  "competences": ["...", "..."],
  "experiences": ["...", "..."],
  "formations": ["...", "..."],
  "autres": ["...", "..."]
}}

Si tu ne trouves pas une section, laisse-la vide, mais structure toujours le JSON comme ci-dessus.

TEXTE DU CV À PARSER :
\"\"\"
{cv_uploaded_text}
\"\"\"
"""
            parsed_cv_json = ask_groq(prompt_parse_cv)
            cv_data = extract_first_json(parsed_cv_json)
            if not cv_data:
                error = "Erreur extraction IA du CV : JSON IA non extrait ou malformé."
                return render_template("index.html", error=error, **context)
            
            # Patch Perso fallback (toujours défini)
            nom = cv_data.get("nom", nom)
            prenom = cv_data.get("prenom", prenom)
            adresse = cv_data.get("adresse", adresse)
            telephone = cv_data.get("telephone", telephone)
            email = cv_data.get("email", email)
            age = cv_data.get("age", age)
            infos_perso = {
                "nom": nom, "prenom": prenom, "adresse": adresse,
                "telephone": telephone, "email": email, "age": age
            }

            prompt_lm_cv = f"""
Voici le parsing structuré du CV du candidat, issu de l’étape précédente :
{json.dumps(cv_data, ensure_ascii=False, indent=2)}

Voici l'offre d'emploi à laquelle il postule :
\"\"\"
{description}
\"\"\"

1. Rédige une lettre de motivation personnalisée et professionnelle adaptée à l'offre et au parcours du candidat (exploite le maximum d’infos utiles, mets en avant les expériences ou compétences transversales si besoin).
2. Génère le contenu d’un CV adapté à l’offre, en sélectionnant :
   - Un paragraphe de profil synthétique (adapté au poste)
   - Les compétences les plus pertinentes (croisées entre CV et offre)
   - Les expériences professionnelles les plus adaptées, sous forme de bullet points (intitulé, entreprise, dates, mission principale)
   - Les formations principales
   - Autres infos utiles

Rends ce JSON strictement :
{{
  "lettre_motivation": "....",
  "cv_adapte": {{
    "profil": "...",
    "competences": ["...", "..."],
    "experiences": ["...", "..."],
    "formations": ["...", "..."],
    "autres": ["...", "..."]
  }}
}}
"""
            result2 = ask_groq(prompt_lm_cv)
            data2 = extract_first_json(result2)
            if not data2:
                error = "Erreur extraction IA LM/CV : JSON IA non extrait ou malformé."
                return render_template("index.html", error=error, **context)

            lettre_motivation = data2.get("lettre_motivation", "")
            cv_adapte = data2.get("cv_adapte", {})

            prompt_fiche_poste = f"""
Lis attentivement l'offre d'emploi suivante et extrait-en les éléments principaux pour générer une fiche de poste structurée, en remplissant strictement ce JSON (sans inventer) :

{{
  "titre": "...",
  "employeur": "...",
  "ville": "...",
  "salaire": "...",
  "type_contrat": "...",
  "missions": ["...", "..."],
  "competences": ["...", "..."],
  "avantages": ["...", "..."],
  "savoir_etre": ["...", "..."],
  "autres": ["..."]
}}

Offre à analyser :
\"\"\"
{description}
\"\"\"
"""
            fiche_poste_json = ask_groq(prompt_fiche_poste)
            fiche_poste = extract_first_json(fiche_poste_json) or {}

            # ---------- Génération fichiers ----------
            cv_pdf_path = os.path.join(TMP_DIR, f"{file_id}_cv.pdf")
            cv_docx_path = os.path.join(TMP_DIR, f"{file_id}_cv.docx")
            lm_pdf_path = os.path.join(TMP_DIR, f"{file_id}_lm.pdf")
            lm_docx_path = os.path.join(TMP_DIR, f"{file_id}_lm.docx")
            fiche_pdf_path = os.path.join(TMP_DIR, f"{file_id}_fiche.pdf")
            fiche_docx_path = os.path.join(TMP_DIR, f"{file_id}_fiche.docx")
            # --- HTML rendering
            with open("templates/cv_template.html", encoding="utf-8") as f:
                cv_html = Template(f.read()).render(cv=cv_adapte, **infos_perso)
            with open("templates/lm_template.html", encoding="utf-8") as f:
                lm_html = Template(f.read()).render(lettre_motivation=lettre_motivation, **infos_perso)
            with open("templates/fiche_poste_template.html", encoding="utf-8") as f:
                fiche_html = Template(f.read()).render(fiche_poste=fiche_poste)
            # --- PDF
            pdfkit.from_string(cv_html, cv_pdf_path, configuration=config)
            pdfkit.from_string(lm_html, lm_pdf_path, configuration=config)
            pdfkit.from_string(fiche_html, fiche_pdf_path, configuration=config)
            # --- DOCX
            render_cv_docx(cv_adapte, infos_perso, cv_docx_path)
            render_lm_docx(lettre_motivation, infos_perso, lm_docx_path)
            render_fiche_docx(fiche_poste, fiche_docx_path)

            return render_template(
                "result.html",
                fiche_poste=fiche_poste,
                cv=cv_adapte,
                lettre_motivation=lettre_motivation,
                infos_perso=infos_perso,
                cv_uploaded_text=cv_uploaded_text,
                cv_pdf=f"{file_id}_cv.pdf",
                cv_docx=f"{file_id}_cv.docx",
                lm_pdf=f"{file_id}_lm.pdf",
                lm_docx=f"{file_id}_lm.docx",
                fiche_pdf=f"{file_id}_fiche.pdf",
                fiche_docx=f"{file_id}_fiche.docx"
            )

        # ------- Pas de CV uploadé, fallback formulaire -------
        if not ((any(x.strip() for x in xp_poste) and any(x.strip() for x in dip_titre)) or description.strip()):
            error = "Veuillez remplir au moins une expérience professionnelle, un diplôme, ou uploader votre CV."
            return render_template("index.html", error=error, **context)

        infos_perso = {
            "nom": nom, "prenom": prenom, "adresse": adresse,
            "telephone": telephone, "email": email, "age": age
        }

        prompt_fields = f"""
Voici les infos saisies par le candidat :

Nom : {nom}
Prénom : {prenom}
Adresse : {adresse}
Téléphone : {telephone}
Email : {email}
Âge : {age}

Expériences professionnelles :
{json.dumps(xp_poste)}
Entreprises : {json.dumps(xp_entreprise)}
Lieux : {json.dumps(xp_lieu)}
Dates début : {json.dumps(xp_debut)}
Dates fin : {json.dumps(xp_fin)}

Diplômes : {json.dumps(dip_titre)}
Lieux : {json.dumps(dip_lieu)}
Dates : {json.dumps(dip_date)}

Voici l'offre d'emploi :
\"\"\"
{description}
\"\"\"

Génère une lettre de motivation adaptée à l’offre et au parcours, puis un CV adapté en JSON :

{{
  "lettre_motivation": "...",
  "cv_adapte": {{
    "profil": "...",
    "competences": ["...", "..."],
    "experiences": ["...", "..."],
    "formations": ["...", "..."],
    "autres": ["...", "..."]
  }}
}}
"""
        result2 = ask_groq(prompt_fields)
        data2 = extract_first_json(result2)
        if not data2:
            error = "Erreur IA ou parsing JSON : JSON IA non extrait ou malformé."
            return render_template("index.html", error=error, **context)
        lettre_motivation = data2.get("lettre_motivation", "")
        cv_adapte = data2.get("cv_adapte", {})

        prompt_fiche_poste = f"""
Lis attentivement l'offre d'emploi suivante et extrait-en les éléments principaux pour générer une fiche de poste structurée, en remplissant strictement ce JSON (sans inventer) :

{{
  "titre": "...",
  "employeur": "...",
  "ville": "...",
  "salaire": "...",
  "type_contrat": "...",
  "missions": ["...", "..."],
  "competences": ["...", "..."],
  "avantages": ["...", "..."],
  "savoir_etre": ["...", "..."],
  "autres": ["..."]
}}

Offre à analyser :
\"\"\"
{description}
\"\"\"
"""
        fiche_poste_json = ask_groq(prompt_fiche_poste)
        fiche_poste = extract_first_json(fiche_poste_json) or {}

        file_id = uuid.uuid4().hex
        # --- HTML rendering
        with open("templates/cv_template.html", encoding="utf-8") as f:
            cv_html = Template(f.read()).render(cv=cv_adapte, **infos_perso)
        with open("templates/lm_template.html", encoding="utf-8") as f:
            lm_html = Template(f.read()).render(lettre_motivation=lettre_motivation, **infos_perso)
        with open("templates/fiche_poste_template.html", encoding="utf-8") as f:
            fiche_html = Template(f.read()).render(fiche_poste=fiche_poste)
        # --- PDF
        cv_pdf_path = os.path.join(TMP_DIR, f"{file_id}_cv.pdf")
        lm_pdf_path = os.path.join(TMP_DIR, f"{file_id}_lm.pdf")
        fiche_pdf_path = os.path.join(TMP_DIR, f"{file_id}_fiche.pdf")
        pdfkit.from_string(cv_html, cv_pdf_path, configuration=config)
        pdfkit.from_string(lm_html, lm_pdf_path, configuration=config)
        pdfkit.from_string(fiche_html, fiche_pdf_path, configuration=config)
        # --- DOCX
        cv_docx_path = os.path.join(TMP_DIR, f"{file_id}_cv.docx")
        lm_docx_path = os.path.join(TMP_DIR, f"{file_id}_lm.docx")
        fiche_docx_path = os.path.join(TMP_DIR, f"{file_id}_fiche.docx")
        render_cv_docx(cv_adapte, infos_perso, cv_docx_path)
        render_lm_docx(lettre_motivation, infos_perso, lm_docx_path)
        render_fiche_docx(fiche_poste, fiche_docx_path)

        return render_template(
            "result.html",
            fiche_poste=fiche_poste,
            cv=cv_adapte,
            lettre_motivation=lettre_motivation,
            infos_perso=infos_perso,
            cv_uploaded_text="",
            cv_pdf=f"{file_id}_cv.pdf",
            cv_docx=f"{file_id}_cv.docx",
            lm_pdf=f"{file_id}_lm.pdf",
            lm_docx=f"{file_id}_lm.docx",
            fiche_pdf=f"{file_id}_fiche.pdf",
            fiche_docx=f"{file_id}_fiche.docx"
        )

    return render_template("index.html", **context)

@app.route('/download/<path:filename>')
def download_file(filename):
    full_path = os.path.join(TMP_DIR, os.path.basename(filename))
    if not os.path.exists(full_path):
        return "Fichier introuvable", 404
    return send_file(full_path, as_attachment=True)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
